import streamlit as st
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

# --- CONFIGURACIÓN DE LA APP ---
st.set_page_config(page_title="SVI - ACTA 10 BIS (Word)", page_icon="👮", layout="wide")

# --- INICIALIZACIÓN DE MEMORIA ---
campos_base = {
    'unidad': "G.T.M.", 'tercio': "TERCIO ALPHA", 'movil': "12.116", 'jurisdiccion': "SUB 18",
    'lugar': "", 'hora_demora': datetime.now().strftime('%H:%M'), 'acta_nro': "",
    'apellido': "", 'nombre': "", 'dni': "", 'nacionalidad': "ARGENTINA",
    'est_civil': "SOLTERO/A", 'edad': "", 'nacimiento': "", 'domicilio': "", 'padres': "",
    'actuante_ap': "", 'actuante_nom': "", 'refuerzo_ap': "", 'refuerzo_nom': "",
    'testigo_datos': "", 'requisa': "PALPADO PREVENTIVO SOBRE SUS PRENDAS NO HALLANDO ELEMENTOS DE PELIGROSIDAD", 
    'secuestro': "PERTENENCIAS PERSONALES", 'estado_fisico': "no presenta lesiones visibles, golpes ni manifiesta dolencias al momento de ingresar a la dependencia",
    'motivo_unico': "", 'of_recibe': "SUBOFICIAL RODRIGUEZ (M)"
}

for clave, valor in campos_base.items():
    if clave not in st.session_state:
        st.session_state[clave] = valor

if 'historial' not in st.session_state:
    st.session_state.historial = []

# --- MOTOR DE GENERACIÓN WORD ---
def generar_word():
    doc = Document()
    
    # Configuración de márgenes (30mm izq para abrochado reglamentario)
    section = doc.sections[0]
    section.left_margin = Mm(30)
    section.right_margin = Mm(20)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)

    # Título central
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titulo.add_run("ACTA DE DEMORA ART 10 BIS LEY 7.395")
    run_t.bold = True
    run_t.font.size = Pt(13)
    run_t.font.name = 'Arial'

    # Cuerpo del Acta
    ahora = datetime.now()
    meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # JUSTIFICADO PERFECTO
    
    def agregar(texto, negrita=False):
        run = p.add_run(texto)
        run.bold = negrita
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    # Redacción técnica con negritas intercaladas
    agregar(f"En la ciudad de ROSARIO, departamento Rosario de la provincia de Santa Fe, a los {ahora.day} días del mes de {meses[ahora.month-1]} del año {ahora.year}, siendo las ")
    agregar(f"{st.session_state.hora_demora} hs", True)
    agregar(", el funcionario policial actuante ")
    agregar(f"{st.session_state.actuante_ap.upper()} {st.session_state.actuante_nom.upper()}", True)
    agregar(f" a cargo de la unidad móvil {st.session_state.movil} juntamente con el refuerzo ")
    agregar(f"{st.session_state.refuerzo_ap.upper()} {st.session_state.refuerzo_nom.upper()}", True)
    agregar(f", ambos pertenecientes a {st.session_state.unidad} de la UR II Rosario, ")
    agregar("SE HACE CONSTAR: ", True)
    agregar("Que de conformidad a lo establecido en el ")
    agregar("Art. 10 bis de la Ley Orgánica de Policial de la Provincia de Santa Fe N° 7395", True)
    agregar(f" se procede a DEMORAR a las {st.session_state.hora_demora} horas, desde calle ")
    agregar(f"{st.session_state.lugar.upper()}", True)
    agregar(" al cual manifiesta ser ")
    agregar(f"{st.session_state.apellido.upper()} {st.session_state.nombre.upper()}, DNI {st.session_state.dni}", True)
    agregar(f", de nacionalidad {st.session_state.nacionalidad.upper()}, estado civil {st.session_state.est_civil}, nacido el {st.session_state.nacimiento}, contando con {st.session_state.edad} años de edad, hijo de {st.session_state.padres.upper()}, con domicilio real en la calle {st.session_state.domicilio.upper()} de esta ciudad. ")
    agregar("Adoptándose esta medida por el motivo de que ante la presencia policial: ")
    agregar(f"{st.session_state.motivo_unico.upper()}. ", True)
    agregar("A continuación se le imponen al demorado sus derechos y garantías: a) será trasladado a la dependencia; ")
    agregar("b) la demora no superará las SEIS (6) HORAS; ", True)
    agregar(f"c) no será incomunicado; d) tiene derecho a realizar una llamada telefónica; e) no será alojado con detenidos comunes; f) se labra la presente acta ante los testigos: {st.session_state.testigo_datos.upper()}. ")
    agregar("Se hace constar que de la requisa efectuada sobre el demorado la misma arrojó resultado ")
    agregar(f"{st.session_state.requisa.upper()}. ", True)
    agregar("Es dable hacer mención que se le secuestra en carácter de depósito para su resguardo: ")
    agregar(f"{st.session_state.secuestro.upper()}. ", True)
    agregar(f"Se deja constancia que el demorado {st.session_state.estado_fisico}. Se hace constar que se labra la presente en recibiendo de conformidad ")
    agregar(f"{st.session_state.of_recibe.upper()}", True)
    agregar(" en carácter de oficial de guardia. Con lo que no siendo para más se da por finalizado el presente acto del cual firman los testigos, demorado y el personal actuante para su debida constancia.")

    doc.add_paragraph("\n\nHORA DE CESE: __________ hs.")

    # Convertir a bytes para descarga
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# --- INTERFAZ ---
st.title("👮 SVI - GENERADOR DE ACTAS 10 BIS")

tab1, tab2, tab3, tab4, tab5 = st.tabs(["🚔 Servicio", "👤 Demorado", "🤝 Testigo/Requisa", "✍️ Cierre y Descarga", "🗂️ Historial"])

with tab1:
    c1, c2 = st.columns(2)
    u_opciones = ["G.T.M.", "C.R.E. ROSARIO", "B.O.U.", "P.A.T.", "OTRO"]
    u_sel = c1.selectbox("Unidad", u_opciones, index=u_opciones.index(st.session_state.unidad) if st.session_state.unidad in u_opciones else 0)
    st.session_state.unidad = c1.text_input("Especifique Unidad", value=st.session_state.unidad) if u_sel == "OTRO" else u_sel
    st.session_state.movil = st.text_input("Móvil / Legajo", value=st.session_state.movil)
    st.session_state.lugar = st.text_input("Lugar del procedimiento", value=st.session_state.lugar)
    st.session_state.hora_demora = st.text_input("Hora de inicio", value=st.session_state.hora_demora)
    st.session_state.acta_nro = st.text_input("Acta Nº", value=st.session_state.acta_nro)

with tab2:
    ca, cb = st.columns(2)
    st.session_state.apellido = ca.text_input("Apellido/s", value=st.session_state.apellido)
    st.session_state.nombre = cb.text_input("Nombre/s", value=st.session_state.nombre)
    st.session_state.dni = st.text_input("DNI", value=st.session_state.dni)
    st.session_state.nacimiento = st.text_input("Fecha de Nacimiento", value=st.session_state.nacimiento)
    st.session_state.domicilio = st.text_input("Domicilio Real", value=st.session_state.domicilio)
    st.session_state.padres = st.text_input("Filiación (Hijo de)", value=st.session_state.padres)

with tab3:
    st.session_state.actuante_ap = st.text_input("Ap. Actuante", value=st.session_state.actuante_ap)
    st.session_state.refuerzo_ap = st.text_input("Ap. Refuerzo", value=st.session_state.refuerzo_ap)
    st.session_state.testigo_datos = st.text_area("Testigos", value=st.session_state.testigo_datos)
    st.session_state.secuestro = st.text_area("Elementos secuestrados", value=st.session_state.secuestro)

with tab4:
    st.session_state.motivo_unico = st.text_area("Redacción del motivo:", value=st.session_state.motivo_unico)
    st.session_state.of_recibe = st.text_input("Oficial de Guardia que recibe", value=st.session_state.of_recibe)

    col_w, col_h = st.columns(2)
    with col_w:
        # BOTÓN GENERAR WORD
        if st.session_state.apellido:
            doc_bytes = generar_word()
            st.download_button(
                label="📝 DESCARGAR ACTA (WORD)",
                data=doc_bytes,
                file_name=f"10Bis_{st.session_state.apellido}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("Cargue el apellido para habilitar la descarga.")

    with col_h:
        if st.button("💾 GUARDAR EN HISTORIAL"):
            reg = {clave: st.session_state[clave] for clave in campos_base.keys()}
            reg['fecha_registro'] = datetime.now().strftime('%H:%M:%S')
            st.session_state.historial.append(reg)
            st.success("Guardado.")

    # --- PARTE DE WHATSAPP ---
    st.subheader("📲 Resumen para WhatsApp")
    res_wa = f"""*🚔 {st.session_state.unidad}* | *ACTA 10 BIS*
*HORA:* {st.session_state.hora_demora} hs.
*👤 DEMORADO:* **{st.session_state.apellido.upper()} {st.session_state.nombre.upper()}**
*DNI:* {st.session_state.dni}
*MOTIVO:* {st.session_state.motivo_unico.upper()}"""
    st.code(res_wa)

with tab5:
    st.subheader("🗂️ Historial")
    for i, reg in enumerate(reversed(st.session_state.historial)):
        if st.button(f"📌 {reg['apellido']} ({reg['fecha_registro']})", key=f"h_{i}"):
            for k in campos_base.keys(): st.session_state[k] = reg[k]
            st.rerun()
